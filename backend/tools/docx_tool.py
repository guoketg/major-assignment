"""
Word 文档生成工具（LangChain Tool 封装）

提供创建和编辑 Word 文档的功能，供 Reporter Agent 使用。
所有文件操作均经过安全护栏检查（路径遍历防护、文件名校验）。
"""
import os
from typing import List, Optional

from langchain_core.tools import tool
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "generated_reports")


def _safe_path(filename: str) -> str:
    """路径安全校验：防止路径遍历攻击"""
    # 禁止空字节、路径遍历字符
    if "\x00" in filename or ".." in filename:
        raise ValueError(f"非法文件名（含路径遍历字符）: {filename[:50]}...")
    # 禁止路径分隔符
    if "/" in filename or "\\" in filename:
        raise ValueError(f"非法文件名（含路径分隔符）: {filename[:50]}...")
    # 文件名长度限制
    if len(filename) > 200:
        raise ValueError(f"文件名过长（最大200字符）: {len(filename)}")
    # 清理特殊字符
    safe_name = "".join(c for c in filename if c.isalnum() or c in " _-().,.，。（）").strip()
    if not safe_name:
        raise ValueError("文件名为空或仅含特殊字符")
    return safe_name


@tool
def create_docx(title: str) -> str:
    """创建一个新的 Word 文档。

    当你需要生成调研报告、实验报告或任何 Word 文档时使用此工具。

    Args:
        title: 文档标题（不含扩展名）

    Returns:
        生成的 .docx 文件路径
    """
    try:
        safe_title = _safe_path(title)
    except ValueError as e:
        return f"[安全护栏拦截] {e}"

    os.makedirs(REPORTS_DIR, exist_ok=True)
    path = os.path.join(REPORTS_DIR, f"{safe_title}.docx")

    # 二次确认路径在允许的目录内
    real_path = os.path.realpath(path)
    if not real_path.startswith(os.path.realpath(REPORTS_DIR)):
        return f"[安全护栏拦截] 文件路径不在允许的目录内"

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # 添加标题
    doc.add_heading(title, 0)

    doc.save(real_path)
    return f"文档已创建: {real_path}"


@tool
def add_section(filepath: str, heading: str, content: str) -> str:
    """在已存在的 Word 文档中添加一个新章节。

    章节包含标题和正文内容。

    Args:
        filepath: 文档路径
        heading: 章节标题
        content: 章节正文（支持 Markdown 风格的换行）

    Returns:
        操作结果
    """
    # 路径安全检查
    if "\x00" in filepath or ".." in filepath or "/" in filepath.replace(os.sep, "/").lstrip(os.sep + "/"):
        return "[安全护栏拦截] 文件路径包含非法字符"

    real_path = os.path.realpath(filepath)
    if not real_path.startswith(os.path.realpath(REPORTS_DIR)):
        return "[安全护栏拦截] 文件路径不在允许的目录内"

    if not os.path.exists(real_path):
        return f"错误: 文件不存在 {filepath}"

    doc = Document(real_path)
    doc.add_heading(heading, level=1)

    for paragraph in content.split("\n\n"):
        p = doc.add_paragraph(paragraph.strip())
        p.paragraph_format.space_after = Pt(6)

    doc.save(real_path)
    return f"章节 '{heading}' 已添加到文档"


@tool
def add_table(filepath: str, headers: List[str], rows: List[List[str]]) -> str:
    """在 Word 文档中添加一个对比表格。

    用于展示方法对比、实验结果等结构化数据。

    Args:
        filepath: 文档路径
        headers: 表头列表，如 ["方法", "准确率", "参数量"]
        rows: 数据行列表，每行是一个列表，长度与 headers 相同

    Returns:
        操作结果
    """
    # 路径安全检查
    if "\x00" in filepath or ".." in filepath or "/" in filepath.replace(os.sep, "/").lstrip(os.sep + "/"):
        return "[安全护栏拦截] 文件路径包含非法字符"

    real_path = os.path.realpath(filepath)
    if not real_path.startswith(os.path.realpath(REPORTS_DIR)):
        return "[安全护栏拦截] 文件路径不在允许的目录内"

    if not os.path.exists(real_path):
        return f"错误: 文件不存在 {filepath}"

    if not headers or not rows:
        return "错误: 表头和数据不能为空"

    doc = Document(real_path)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    # 表头
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    # 数据行
    for r, row in enumerate(rows):
        if len(row) != len(headers):
            continue
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = str(val)

    doc.save(real_path)
    return f"表格 ({len(headers)}列 x {len(rows)}行) 已添加到文档"


# 工具列表
docx_tools = [create_docx, add_section, add_table]
