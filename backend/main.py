"""
多轮对话API服务

基于FastAPI构建，支持多轮对话、历史记录管理、arXiv论文搜索、Agent可视化
"""
import os
import json
import time
import uuid
import asyncio
import glob
import logging
from datetime import datetime
from functools import lru_cache
from typing import Optional

from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel, Field
from openai import OpenAI
import httpx
import xmltodict
from backend.agent.graph import AgentGraph
from backend.agent.guardrails import get_guardrail
from backend.agent.llmops import get_llmops

load_dotenv()

# 日志配置
logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
logger = logging.getLogger(__name__)


def sanitize_log(text: str, max_len: int = 60) -> str:
    """日志脱敏：截断过长内容，替换换行符"""
    if not text:
        return ""
    s = text.replace('\n', ' ').replace('\r', ' ').strip()
    return s[:max_len] + '...' if len(s) > max_len else s

# ====== 对话持久化存储 ======

CONVERSATIONS_DIR = "conversations"


def _ensure_conv_dir():
    os.makedirs(CONVERSATIONS_DIR, exist_ok=True)


def _conv_path(session_id: str) -> str:
    return os.path.join(CONVERSATIONS_DIR, f"{session_id}.json")


def _load_all_conversations():
    """启动时从文件加载所有历史对话"""
    _ensure_conv_dir()
    loaded = 0
    for path in glob.glob(os.path.join(CONVERSATIONS_DIR, "*.json")):
        # 跳过 meta 文件
        if path.endswith(".meta.json"):
            continue
        sid = os.path.basename(path).replace(".json", "")
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            # 兼容新旧格式
            if isinstance(data, list):
                conversations[sid] = data
            elif isinstance(data, dict) and "history" in data:
                conversations[sid] = data
            else:
                continue
            loaded += 1
        except (json.JSONDecodeError, IOError):
            pass
    if loaded:
        print(f"  已加载 {loaded} 个历史会话")


def _save_conversation(session_id: str):
    """保存单个会话到文件（统一使用列表格式）"""
    if session_id in conversations:
        _ensure_conv_dir()
        path = _conv_path(session_id)
        try:
            conv = conversations[session_id]
            # 统一以列表格式保存（兼容 dict 格式降级）
            if isinstance(conv, list):
                data = conv
            elif isinstance(conv, dict):
                data = conv.get("history", [])
            else:
                return
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except IOError:
            pass


def _delete_conversation_file(session_id: str):
    """删除会话文件及 meta 文件"""
    for suffix in ["", ".meta"]:
        path = os.path.join(CONVERSATIONS_DIR, f"{session_id}{suffix}.json")
        if os.path.exists(path):
            try:
                os.remove(path)
            except IOError:
                pass


def _pipeline_meta_path(session_id: str) -> str:
    """流水线 meta 文件路径"""
    return os.path.join(CONVERSATIONS_DIR, f"{session_id}.meta.json")


def _save_pipeline_meta(session_id: str, data: dict):
    """保存流水线 meta 数据到单独文件"""
    _ensure_conv_dir()
    path = _pipeline_meta_path(session_id)
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except IOError:
        pass


def _load_pipeline_meta(session_id: str) -> dict:
    """加载流水线 meta 数据"""
    path = _pipeline_meta_path(session_id)
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            pass
    return {}


def _get_history(session_id: str) -> list:
    """安全获取会话历史列表（兼容新旧格式）"""
    conv = conversations.get(session_id, [])
    if isinstance(conv, list):
        return conv
    if isinstance(conv, dict):
        return conv.get("history", [])
    return []


# ====== FastAPI 应用 ======

app = FastAPI(title="多轮对话服务", version="1.0.0")


@app.on_event("startup")
async def startup():
    """启动时加载历史会话"""
    _load_all_conversations()


# CORS配置
# CORS配置 — 限制为已知前端地址
CORS_ORIGINS = os.getenv("CORS_ORIGINS", "http://localhost:3000,http://localhost:3001,http://127.0.0.1:3000")
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS.split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# OpenAI客户端配置
api_key = os.getenv("DEEPSEEK_API_KEY") or os.getenv("OPENAI_API_KEY", "")
DEFAULT_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")
client = OpenAI(
    api_key=api_key,
    base_url=os.getenv("OPENAI_BASE_URL", "https://api.deepseek.com/v1")
)

# 会话存储 (生产环境应使用数据库)
conversations: dict[str, list | dict] = {}

# 会话预算限制
MAX_SESSIONS = 200  # 最大内存会话数
MAX_TOKENS_PER_SESSION = 200000  # 每会话最大 Token 数（约 ¥0.3）
MAX_TURNS_PER_SESSION = 100     # 每会话最大轮次

# 模型映射（前端显示名 -> API 模型名）
MODEL_MAP = {
    "deepseek-v4-flash": "deepseek-chat",
    "v4-pro": "deepseek-chat",
    "思考模式": "deepseek-reasoner",
}

# Agent Graph 实例（懒加载）
_agent_graph = None


def get_agent_graph() -> AgentGraph:
    """获取或创建全局 AgentGraph 实例"""
    global _agent_graph
    if _agent_graph is None:
        _agent_graph = AgentGraph()
    return _agent_graph


class ChatRequest(BaseModel):
    session_id: str = Field(..., min_length=1, max_length=64)
    message: str = Field(..., min_length=1, max_length=10000, description="用户消息，1-10000字符")
    model: str = ""  # 可选，前端选择的模型名
    agent: str = Field(default="auto", pattern=r"^(auto|chat|research|innovate|experiment)$")  # 可选
    skill: str = Field(default="none", max_length=32)  # 技能增强模式（支持自定义技能）


class ChatResponse(BaseModel):
    session_id: str
    reply: str
    history: list[dict]


@app.post("/chat/stream")
async def chat_stream(request: ChatRequest):
    """流式处理用户消息，返回AI回复"""
    session_id = request.session_id
    user_message = request.message

    # 初始化会话历史（兼容新旧格式）
    if session_id not in conversations:
        conversations[session_id] = []
    history_list = _get_history(session_id)

    # 添加用户消息到历史（立即保存）
    history_list.append({
        "role": "user",
        "content": user_message,
        "timestamp": datetime.now().isoformat()
    })
    conversations[session_id] = history_list
    _save_conversation(session_id)

    # 确定使用的模型
    selected_model = MODEL_MAP.get(request.model, "") or os.getenv("DEEPSEEK_MODEL", "deepseek-chat")

    async def generate():
        nonlocal selected_model, user_message
        try:
            # === 预算限制检查：轮次上限 ===
            history_list_check = _get_history(session_id)
            user_turns = sum(1 for m in history_list_check if isinstance(m, dict) and m.get("role") == "user")
            if user_turns > MAX_TURNS_PER_SESSION:
                yield f"data: {json.dumps({'type': 'error', 'error': f'对话已达最大轮次限制（{MAX_TURNS_PER_SESSION}轮），请创建新会话', 'done': True})}\n\n"
                return

            # === 安全护栏：输入检查 ===
            guardrail = get_guardrail()
            input_check = guardrail.check_input(user_message, session_id)
            if not input_check.passed:
                blocked_msg = f"[安全护栏] {input_check.message}"
                logger.warning(f"[GUARDRAIL_BLOCK] session={sanitize_log(session_id)}, reason={input_check.reason.name}, msg={input_check.message[:80]}")
                # 记录拦截到会话历史
                history_list.append({
                    "role": "assistant",
                    "content": blocked_msg,
                    "timestamp": datetime.now().isoformat(),
                    "guardrail_blocked": True,
                })
                conversations[session_id] = history_list
                _save_conversation(session_id)
                yield f"data: {json.dumps({'type': 'error', 'error': blocked_msg, 'done': True, 'guardrail_blocked': True})}\n\n"
                return
            # 输入通过护栏后，如果内容被脱敏则使用脱敏后的内容
            if input_check.sanitized_content:
                user_message = input_check.sanitized_content
                # 更新历史记录中最后一条用户消息
                history_list[-1]["content"] = user_message
                history_list[-1]["sanitized"] = True

            # === 使用 LangGraph 多 Agent 引擎 ===
            agent_graph = get_agent_graph()

            async with asyncio.timeout(600):
                async for event in agent_graph.run_stream(
                    session_id=session_id,
                    messages=_get_history(session_id),
                    model=selected_model,
                    agent=request.agent,
                    skill=request.skill,
                ):
                    # 直接转发 AgentGraph 的事件到 SSE
                    yield f"data: {json.dumps(event)}\n\n"

                    # 如果收到 done 或 error 事件，保存历史
                    if event.get("done") and event.get("history"):
                        history = event["history"]
                        # 检查 history 的完整性
                        if history and len(history) > 0:
                            last = history[-1]
                            if last.get("role") == "assistant" and not last.get("content"):
                                logger.error(f"[CRITICAL] Assistant content empty in done event! session={sanitize_log(session_id)}, history_len={len(history)}")
                            else:
                                logger.info(f"[SAVE] session={sanitize_log(session_id)}, history_len={len(history)}, last_role={last.get('role','?')}, last_content_len={len(last.get('content',''))}")
                        else:
                            logger.warning(f"[SAVE] Empty history for session={sanitize_log(session_id)}")
                        conversations[session_id] = history
                        _save_conversation(session_id)
                        # 保存流水线 meta 数据到单独文件
                        meta = {
                            "agent_pipeline": event.get("agent_pipeline", []),
                            "tool_calls": event.get("tool_calls", []),
                            "sub_task_plan": event.get("sub_task_plan", []),
                            "token_usage": event.get("token_usage", {}),
                            "total_cost": event.get("total_cost", 0.0),
                            "per_agent_tokens": event.get("per_agent_tokens", {}),
                        }
                        _save_pipeline_meta(session_id, meta)

                        # Token 预算警告
                        token_usage = event.get("token_usage", {})
                        total_tokens = token_usage.get("total_tokens", 0) if isinstance(token_usage, dict) else 0
                        if total_tokens > MAX_TOKENS_PER_SESSION * 0.8:
                            warning = f"⚠️ 本次会话已使用约 {total_tokens} Token（上限 {MAX_TOKENS_PER_SESSION}），请注意控制对话长度"
                            yield f"data: {json.dumps({'type': 'warning', 'content': warning})}\n\n"

        except asyncio.TimeoutError:
            yield f"data: {json.dumps({'type': 'error', 'error': '执行超时（600秒），请简化您的请求或稍后重试', 'done': True})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'error': str(e), 'done': True})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.post("/chat", response_model=ChatResponse)
async def chat_non_stream(request: ChatRequest):
    """处理用户消息，返回AI回复（非流式）"""
    session_id = request.session_id
    user_message = request.message

    # 初始化会话历史
    if session_id not in conversations:
        conversations[session_id] = []
    history_list = _get_history(session_id)

    # 添加用户消息到历史（立即保存）
    history_list.append({
        "role": "user",
        "content": user_message,
        "timestamp": datetime.now().isoformat()
    })
    conversations[session_id] = history_list
    _save_conversation(session_id)

    # === 安全护栏：输入检查 ===
    guardrail = get_guardrail()
    input_check = guardrail.check_input(user_message, session_id)
    if not input_check.passed:
        raise HTTPException(status_code=400, detail=f"[安全护栏] {input_check.message}")

    try:
        # 调用DeepSeek API
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=_get_history(session_id)
        )

        assistant_reply = response.choices[0].message.content

        # 添加助手回复到历史
        history_list = _get_history(session_id)
        history_list.append({
            "role": "assistant",
            "content": assistant_reply,
            "timestamp": datetime.now().isoformat()
        })
        conversations[session_id] = history_list
        _save_conversation(session_id)

        return ChatResponse(
            session_id=session_id,
            reply=assistant_reply,
            history=_get_history(session_id)
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"API调用失败: {str(e)}")


@app.get("/history/{session_id}")
async def get_history(session_id: str):
    """获取会话历史（含流水线 meta 数据）"""
    if session_id not in conversations:
        return {"session_id": session_id, "history": []}

    history = _get_history(session_id)
    meta = _load_pipeline_meta(session_id)

    # 旧会话兼容：有消息但没有 meta 文件时自动生成基础流水线
    if not meta and history:
        has_assistant = any(isinstance(m, dict) and m.get("role") == "assistant" for m in history)
        if has_assistant:
            meta = {
                "agent_pipeline": [
                    {"agent": "supervisor", "label": "🤖 智能路由", "status": "complete"},
                    {"agent": "chat_agent", "label": "💬 对话助手", "status": "complete"},
                    {"agent": "synthesizer", "label": "📋 综合输出", "status": "complete"},
                ],
                "tool_calls": [],
                "sub_task_plan": [],
            }
            _save_pipeline_meta(session_id, meta)

    return {
        "session_id": session_id,
        "history": history,
        "meta": meta,
    }


@app.get("/sessions")
async def list_sessions():
    """获取所有会话列表"""
    session_list = []
    for session_id, conv in conversations.items():
        history = _get_history(session_id)
        if not history:
            continue
        first_user_msg = next(
            (msg["content"][:30] + "..." if len(msg["content"]) > 30 else msg["content"]
             for msg in history if isinstance(msg, dict) and msg.get("role") == "user"),
            "新会话"
        )
        session_list.append({
            "session_id": session_id,
            "title": first_user_msg,
            "message_count": len(history),
            "last_update": history[-1]["timestamp"] if history and isinstance(history[-1], dict) else None
        })
    session_list.sort(key=lambda x: x["last_update"] or "", reverse=True)
    return {"sessions": session_list}


@app.post("/sessions")
async def create_session():
    """创建新会话（含会话数量上限检查）"""
    # 检查会话数量上限
    if len(conversations) >= MAX_SESSIONS:
        # 删除最早使用的会话
        oldest_sid = None
        oldest_ts = None
        for sid in list(conversations.keys()):
            history = _get_history(sid)
            if history and isinstance(history[-1], dict):
                ts = history[-1].get("timestamp", "")
                if oldest_ts is None or ts < oldest_ts:
                    oldest_ts = ts
                    oldest_sid = sid
        if oldest_sid:
            del conversations[oldest_sid]
            _delete_conversation_file(oldest_sid)

    new_session_id = f"session_{uuid.uuid4().hex}"
    conversations[new_session_id] = []
    _save_conversation(new_session_id)
    return {"session_id": new_session_id}


@app.delete("/history/{session_id}")
async def delete_history(session_id: str):
    """删除会话历史"""
    if session_id in conversations:
        del conversations[session_id]
    _delete_conversation_file(session_id)
    return {"message": "会话已删除"}


@app.get("/health")
async def health_check():
    """健康检查"""
    return {"status": "ok", "time": datetime.now().isoformat()}


# ====== LLMOps 可观测性 API ======


@app.get("/llmops/health")
async def llmops_health():
    """LLMOps 健康状态检查"""
    llmops = get_llmops()
    snap = llmops.metrics.snapshot()
    return {
        "status": "ok",
        "time": datetime.now().isoformat(),
        "cache_size": llmops.cache.size,
        "total_calls": snap.total_calls,
        "error_rate": round(snap.error_calls / max(snap.total_calls, 1), 4),
        "latency_p95_ms": snap.latency_p95_ms,
        "rate_limiter_tokens": round(llmops.rate_limiter.available_tokens, 1),
    }


@app.get("/llmops/metrics")
async def llmops_metrics():
    """获取 LLMOps 性能指标快照（延迟、成功率、Token 吞吐量等）"""
    llmops = get_llmops()
    return llmops.metrics.snapshot().to_dict()


@app.get("/llmops/traces")
async def llmops_traces(limit: int = Query(100, ge=1, le=500)):
    """获取最近 LLM 调用追踪记录"""
    llmops = get_llmops()
    traces = llmops.tracer.get_recent_traces(limit=limit)
    # 同时附加按 agent 的聚合统计
    per_agent = llmops.metrics.get_per_agent_snapshot(traces)
    return {"total": len(traces), "traces": traces, "per_agent": per_agent}


@app.get("/llmops/traces/{session_id}")
async def llmops_session_traces(session_id: str, limit: int = Query(50, ge=1, le=200)):
    """按会话 ID 查询 LLM 调用追踪记录"""
    llmops = get_llmops()
    traces = llmops.tracer.get_traces_by_session(session_id, limit=limit)
    return {"session_id": session_id, "total": len(traces), "traces": traces}


@app.get("/llmops/alerts")
async def llmops_alerts(
    limit: int = Query(50, ge=1, le=200),
    acknowledged: Optional[bool] = Query(None),
):
    """获取 LLMOps 告警列表"""
    llmops = get_llmops()
    alerts = llmops.alerts.get_alerts(limit=limit, acknowledged=acknowledged)
    return {"total": len(alerts), "alerts": alerts}


@app.post("/llmops/alerts/{alert_id}/acknowledge")
async def llmops_acknowledge_alert(alert_id: str):
    """确认（标记已读）指定告警"""
    llmops = get_llmops()
    success = llmops.alerts.acknowledge(alert_id)
    if not success:
        raise HTTPException(status_code=404, detail="告警不存在")
    return {"message": "已确认", "alert_id": alert_id}


@app.get("/llmops/cache")
async def llmops_cache_stats():
    """获取 LLM 响应缓存统计"""
    llmops = get_llmops()
    return {
        **llmops.cache.stats,
        "hit_rate": round(
            llmops.metrics.snapshot().cache_hits / max(
                llmops.metrics.snapshot().cache_hits + llmops.metrics.snapshot().cache_misses, 1
            ), 4
        ),
    }


@app.post("/llmops/cache/clear")
async def llmops_cache_clear():
    """清空 LLM 响应缓存"""
    llmops = get_llmops()
    llmops.cache.clear()
    return {"message": "缓存已清空"}


@app.post("/llmops/metrics/persist")
async def llmops_metrics_persist():
    """手动持久化当日指标汇总"""
    llmops = get_llmops()
    llmops.tracer.flush()
    llmops.metrics.persist_daily_summary()
    return {"message": "指标已持久化", "snapshot": llmops.metrics.snapshot().to_dict()}


# ====== Word 文档导出 ======


@app.get("/export/docx/{session_id}")
async def export_session_docx(session_id: str):
    """将会话历史导出为 Word 文档"""
    from backend.tools.docx_tool import REPORTS_DIR

    if session_id not in conversations:
        raise HTTPException(status_code=404, detail="会话不存在")

    history = _get_history(session_id)
    if not history:
        raise HTTPException(status_code=400, detail="会话没有消息")

    # 获取第一条用户消息作为文档标题
    title = next(
        (msg["content"][:20] + "..." if len(msg["content"]) > 20 else msg["content"]
         for msg in history if isinstance(msg, dict) and msg.get("role") == "user"),
        f"对话记录_{session_id}"
    )
    safe_title = "".join(c for c in title if c.isalnum() or c in " _-，。").strip()

    os.makedirs(REPORTS_DIR, exist_ok=True)
    filename = f"{safe_title}_{session_id}.docx"
    filepath = os.path.join(REPORTS_DIR, filename)

    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # 标题
    doc.add_heading(f"AI 对话记录: {title}", 0)
    doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")

    # 遍历消息
    for msg in history:
        role = msg.get("role", "unknown")
        content = msg.get("content", "")
        ts = msg.get("timestamp", "")

        if role == "user":
            p = doc.add_paragraph()
            run = p.add_run(f"👤 用户")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x00, 0x7B, 0xFF)
            if ts:
                p.add_run(f"  ({ts[:19]})").font.size = Pt(9)
            doc.add_paragraph(content)
        elif role == "assistant":
            p = doc.add_paragraph()
            run = p.add_run(f"🤖 AI")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)
            if ts:
                p.add_run(f"  ({ts[:19]})").font.size = Pt(9)

            if msg.get("reasoning_content"):
                doc.add_paragraph("【思考过程】")
                doc.add_paragraph(msg["reasoning_content"])

            doc.add_paragraph(content)
        else:
            doc.add_paragraph(f"[{role}]: {content}")

        doc.add_paragraph("")

    doc.save(filepath)

    return FileResponse(
        filepath,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )


@app.get("/download/reports/{filename}")
async def download_report(filename: str):
    """下载生成的 Word 报告文档"""
    # 安全校验：禁止路径遍历、null字符、非docx文件
    if "\x00" in filename or ".." in filename or "/" in filename or "\\" in filename:
        raise HTTPException(status_code=400, detail="非法文件名")
    if not filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="仅支持 .docx 文件")
    if len(filename) > 200:
        raise HTTPException(status_code=400, detail="文件名过长")

    from backend.tools.docx_tool import REPORTS_DIR
    filepath = os.path.join(REPORTS_DIR, filename)
    real_path = os.path.realpath(filepath)
    if not real_path.startswith(os.path.realpath(REPORTS_DIR)):
        raise HTTPException(status_code=403, detail="非法路径")
    if not os.path.exists(real_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(
        real_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )


@app.get("/memory/{session_id}")
async def get_memory(session_id: str):
    """获取会话的工作记忆（论文归档、创新方案、实验日志）"""
    from backend.memory.manager import load_working_memory
    memory = load_working_memory(session_id)
    history = _get_history(session_id)
    return {
        "session_id": session_id,
        "memory": memory,
        "has_history": len(history) > 0,
        "stats": {
            "papers": len(memory.get("papers_archive", [])),
            "innovations": len(memory.get("innovation_candidates", [])),
            "experiments": len(memory.get("experiment_log", [])),
        },
    }


# ====== arXiv 论文搜索 ======

ARXIV_API_URL = "https://export.arxiv.org/api/query"
_last_arxiv_req_time = 0


@app.get("/arxiv/search")
async def arxiv_search(
    q: str = Query(..., description="搜索查询"),
    start: int = Query(0, ge=0),
    max_results: int = Query(10, ge=1, le=100),
    sortBy: str = Query("relevance", pattern="^(relevance|submittedDate|lastUpdatedDate)$"),
    sortOrder: str = Query("descending", pattern="^(ascending|descending)$"),
):
    """搜索arXiv论文"""
    global _last_arxiv_req_time

    now = time.time()
    since_last = now - _last_arxiv_req_time
    if since_last < 3:
        wait = 3 - since_last
        time.sleep(wait)

    params = {
        "search_query": q,
        "start": start,
        "max_results": max_results,
        "sortBy": sortBy,
        "sortOrder": sortOrder,
    }

    try:
        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.get(ARXIV_API_URL, params=params)
            resp.raise_for_status()
            _last_arxiv_req_time = time.time()

            data = xmltodict.parse(resp.text)
            feed = data.get("feed", {})
            entries = feed.get("entry", [])
            if isinstance(entries, dict):
                entries = [entries]

            papers = []
            for entry in entries:
                authors = entry.get("author", [])
                if isinstance(authors, dict):
                    authors = [authors]

                paper = {
                    "id": entry.get("id", ""),
                    "title": entry.get("title", "").strip().replace("\n", " "),
                    "summary": entry.get("summary", "").strip().replace("\n", " "),
                    "authors": [a.get("name", "") for a in authors],
                    "published": entry.get("published", ""),
                    "updated": entry.get("updated", ""),
                    "categories": [],
                    "pdf_link": "",
                    "links": [],
                }

                cats = entry.get("category", [])
                if isinstance(cats, dict):
                    cats = [cats]
                paper["categories"] = [c.get("@term", "") for c in cats]

                links = entry.get("link", [])
                if isinstance(links, dict):
                    links = [links]
                for link in links:
                    href = link.get("@href", "")
                    paper["links"].append(href)
                    if href.endswith("pdf"):
                        paper["pdf_link"] = href

                papers.append(paper)

            return {
                "total_results": int(feed.get("opensearch:totalResults", 0)),
                "start_index": int(feed.get("opensearch:startIndex", 0)),
                "items_per_page": int(feed.get("opensearch:itemsPerPage", 0)),
                "papers": papers,
            }

    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="arXiv请求超时，请稍后重试")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"arXiv搜索失败: {str(e)}")


# ====== 全局 Token 用量 ======

@app.get("/usage/total")
async def get_total_usage_api():
    """获取全局累计 Token 用量"""
    from backend.agent.total_usage import get_total_usage
    return get_total_usage()


@app.get("/usage/daily")
async def get_daily_usage_api():
    """获取每日 Token 用量列表（按日期倒序）"""
    from backend.agent.total_usage import get_daily_usage
    return get_daily_usage()


@app.get("/usage/date/{date}")
async def get_usage_by_date_api(date: str):
    """获取指定日期的 Token 用量 (YYYY-MM-DD)"""
    from backend.agent.total_usage import get_usage_for_date
    return get_usage_for_date(date)


# ====== 技能管理 API ======

class SkillCreateRequest(BaseModel):
    id: str = Field(..., min_length=1, max_length=32, pattern=r"^[a-z0-9_]+$")
    label: str = Field(..., min_length=1, max_length=32)
    desc: str = Field(default="", max_length=200)
    system_prompt_append: str = Field(default="", max_length=5000)


class SkillUpdateRequest(BaseModel):
    label: str = Field(..., min_length=1, max_length=32)
    desc: str = Field(default="", max_length=200)
    system_prompt_append: str = Field(default="", max_length=5000)


@app.get("/skills")
async def list_skills():
    """获取所有可用技能列表"""
    from backend.agent.skills_store import get_all_skills
    return {"skills": get_all_skills()}


@app.get("/skills/{skill_id}")
async def get_skill(skill_id: str):
    """获取某个技能的详细信息（含 prompt）"""
    from backend.agent.skills_store import get_skill_detail
    detail = get_skill_detail(skill_id)
    if detail is None:
        raise HTTPException(status_code=404, detail="技能不存在")
    return detail


@app.post("/skills")
async def create_skill(req: SkillCreateRequest):
    """创建自定义技能"""
    from backend.agent.skills_store import create_skill as cs
    try:
        return cs(req.id, req.label, req.desc, req.system_prompt_append)
    except ValueError as e:
        raise HTTPException(status_code=409, detail=str(e))


@app.put("/skills/{skill_id}")
async def update_skill(skill_id: str, req: SkillUpdateRequest):
    """更新自定义技能"""
    from backend.agent.skills_store import update_skill as us
    try:
        return us(skill_id, req.label, req.desc, req.system_prompt_append)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.delete("/skills/{skill_id}")
async def delete_skill(skill_id: str):
    """删除自定义技能"""
    from backend.agent.skills_store import delete_skill as ds
    try:
        ds(skill_id)
        return {"message": "已删除"}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
