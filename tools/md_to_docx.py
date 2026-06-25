"""
将综合实验报告 Markdown 转换为 Word (.docx) 格式

用法: python tools/md_to_docx.py
"""
import re
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None, font_name=None):
    """添加格式化 run"""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if size:
        run.font.size = Pt(size)
    if font_name:
        run.font.name = font_name
        # 设置中文字体
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = __import__('lxml.etree', fromlist=['etree']).SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:eastAsia'), font_name)
    return run


def convert_md_to_docx(md_path, docx_path):
    """将 Markdown 文件转换为 Word 文档"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_after = Pt(4)

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    in_table = False
    in_list = False

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                i += 1
                continue
            else:
                in_code_block = False
                i += 1
                continue

        if in_code_block:
            # 代码块内容用等宽字体
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            add_run(p, line, font_name='Courier New', size=9)
            p.style = doc.styles['Normal']
            # 灰色背景
            shading = __import__('lxml.etree', fromlist=['etree']).SubElement(
                p._element.get_or_add_pPr(), qn('w:shd'))
            shading.set(qn('w:fill'), 'F0F0F0')
            shading.set(qn('w:val'), 'clear')
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 分隔线
        if re.match(r'^---+\s*$', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            add_run(p, '─' * 60, color=(180, 180, 180))
            i += 1
            continue

        # 标题（先处理 h1-h6）
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()

            # 清理 Markdown 格式标记
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)

            if level == 1:
                heading = doc.add_heading(text, level=1)
            elif level == 2:
                heading = doc.add_heading(text, level=2)
            elif level == 3:
                heading = doc.add_heading(text, level=3)
            else:
                heading = doc.add_heading(text, level=4)

            # 调整标题字体
            for run in heading.runs:
                run.font.name = 'Arial'
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = __import__('lxml.etree', fromlist=['etree']).SubElement(rPr, qn('w:rFonts'))
                rFonts.set(qn('w:eastAsia'), '微软雅黑')

            i += 1
            continue

        # 表格（以 | 开头和结尾的行）
        if line.strip().startswith('|') and line.strip().endswith('|'):
            # 收集表格所有行
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                if not re.match(r'^\|[\s\-:|]+\|$', lines[i]):  # 跳过分隔行
                    table_rows.append(lines[i])
                i += 1

            if table_rows:
                # 解析表格
                data = []
                for row_line in table_rows:
                    cells = [c.strip() for c in row_line.split('|')[1:-1]]
                    data.append(cells)

                if len(data) >= 2:
                    headers = data[0]
                    rows_data = data[1:]

                    # 确定列数
                    num_cols = max(len(h) for h in data)

                    table = doc.add_table(rows=len(data), cols=num_cols)
                    table.style = 'Light Grid Accent 1'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # 设置表格宽度
                    for cell in table.columns[0].cells:
                        cell.width = Cm(3)

                    # 填入数据
                    for r_idx, row_data in enumerate(data):
                        for c_idx, cell_text in enumerate(row_data):
                            if c_idx < num_cols:
                                cell = table.cell(r_idx, c_idx)
                                cell.text = ''
                                p = cell.paragraphs[0]
                                is_bold = (r_idx == 0)

                                # 处理加粗标记 **text**
                                parts = re.split(r'(\*\*.+?\*\*)', cell_text)
                                for part in parts:
                                    if part.startswith('**') and part.endswith('**'):
                                        add_run(p, part[2:-2], bold=True, size=9)
                                    else:
                                        add_run(p, part, size=9)

                    doc.add_paragraph()  # 表后留空
            continue

        # 列表项（- 或 * 开头）
        list_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(2)

            # 检查是否是任务列表 [x] 或 [ ]
            task_match = re.match(r'^[-*]\s+\[([ x])\]\s+(.+)$', line)

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            # 处理加粗 **text**
            parts = re.split(r'(\*\*.+?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    add_run(p, '• ' + part[2:-2], bold=True)
                else:
                    prefix = '• ' if parts.index(part) == 0 else ''
                    add_run(p, prefix + part)

            i += 1
            continue

        # 有序列表（1. 2. 开头）
        ol_match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
        if ol_match:
            indent = len(ol_match.group(1))
            num = ol_match.group(2)
            text = ol_match.group(3)

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            add_run(p, f'{num}. ')
            # 处理加粗
            parts = re.split(r'(\*\*.+?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    add_run(p, part[2:-2], bold=True)
                else:
                    add_run(p, part)

            i += 1
            continue

        # 普通段落（包含 Markdown 内联格式）
        p = doc.add_paragraph()
        text = line.strip()

        # 处理内联格式：**bold**, *italic*, `code`, [link](url)
        # 首先处理加粗 **text**
        parts = re.split(r'(\*\*.+?\*\*|`[^`]+`|\[.+?\]\(.+?\))', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                add_run(p, part[2:-2], bold=True)
            elif part.startswith('`') and part.endswith('`'):
                add_run(p, part[1:-1], font_name='Courier New', size=9)
            elif part.startswith('[') and '](' in part:
                # 链接
                link_match = re.match(r'\[(.+?)\]\((.+?)\)', part)
                if link_match:
                    add_run(p, link_match.group(1), color=(0, 100, 200))
                    add_run(p, f' ({link_match.group(2)})', size=8, color=(100, 100, 100))
            else:
                # 处理斜体
                italic_parts = re.split(r'(\*[^*]+\*)', part)
                for ip in italic_parts:
                    if ip.startswith('*') and ip.endswith('*') and not ip.startswith('**'):
                        add_run(p, ip[1:-1], italic=True)
                    else:
                        add_run(p, ip)

        i += 1

    # 保存
    doc.save(docx_path)
    print(f'✅ 报告已生成: {docx_path}')
    return docx_path


if __name__ == '__main__':
    md_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                           'docs', '综合实验报告.md')
    docx_path = md_path.replace('.md', '.docx')
    convert_md_to_docx(md_path, docx_path)
